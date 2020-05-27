# -*- coding: utf-8 -*-
"""
Created on Sat May  2 18:37:34 2020

@author: cvicentm
"""

from docx import Document
from docx.shared import RGBColor
from nltk import word_tokenize



def wordToDocument (word, color_hex):
    color_rgb = ImageColor.getrgb(hex)
    run=paragraph.add_run(word+" ")
    font = run.font
    font.color.rgb = RGBColor(color_rgb[0], color_rgb[1], color_rgb[2])

texto = 'En un lugar de la mancha de cuyo nombre no quiero acordarme'
texto2 = 'Vivía un hidalgo de los de lanza en astillero'
tokenize_text = word_tokenize(texto)
tokenize_text2 = word_tokenize(texto2)

document = Document()
paragraph = document.add_paragraph()
[wordToDocument(elemento,0,0,254) for elemento in tokenize_text]
[wordToDocument(elemento,0,254,0) for elemento in tokenize_text2]

document.save('demo1.docx')